#!/usr/bin/env python3
"""
Script para gerar documento Word com cenários de teste - CelestaSupply
Uso: python generate_test_scenarios.py
Saída: test-scenarios.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Dados dos cenários de teste
TESTE_SCENARIOS = [
    # AUTENTICAÇÃO (5)
    {
        "id": "T-001",
        "modulo": "AUTENTICAÇÃO",
        "descricao": "Login com credenciais válidas",
        "pre_condicoes": "Usuário ativo cadastrado",
        "passos": "1. Acessar /login\n2. Preencher email e senha\n3. Clicar 'Entrar'",
        "resultado_esperado": "Redireciona para /dashboard com toast sucesso",
    },
    {
        "id": "T-002",
        "modulo": "AUTENTICAÇÃO",
        "descricao": "Login com email inválido",
        "pre_condicoes": "Sistema online",
        "passos": "1. Acessar /login\n2. Preencher email inexistente\n3. Clicar 'Entrar'",
        "resultado_esperado": "Mensagem de erro 'Email não encontrado'",
    },
    {
        "id": "T-003",
        "modulo": "AUTENTICAÇÃO",
        "descricao": "Login com senha incorreta",
        "pre_condicoes": "Usuário ativo com email válido",
        "passos": "1. Acessar /login\n2. Preencher email correto e senha errada\n3. Clicar 'Entrar'",
        "resultado_esperado": "Mensagem de erro 'Senha inválida'",
    },
    {
        "id": "T-004",
        "modulo": "AUTENTICAÇÃO",
        "descricao": "Login com usuário inativo",
        "pre_condicoes": "Usuário com isActive=false",
        "passos": "1. Acessar /login\n2. Preencher credenciais do usuário inativo\n3. Clicar 'Entrar'",
        "resultado_esperado": "Mensagem de erro 'Usuário inativo ou não autorizado'",
    },
    {
        "id": "T-005",
        "modulo": "AUTENTICAÇÃO",
        "descricao": "Logout",
        "pre_condicoes": "Usuário autenticado",
        "passos": "1. Clicar em menu de usuário\n2. Clicar 'Logout'",
        "resultado_esperado": "Redireciona para /login, sessão encerrada",
    },
    # USUÁRIOS (8)
    {
        "id": "T-006",
        "modulo": "USUÁRIOS",
        "descricao": "Admin cria novo usuário Requester",
        "pre_condicoes": "Admin autenticado, formulário aberto",
        "passos": "1. Preencher: name, email, password, role=requester\n2. Clicar 'Salvar'",
        "resultado_esperado": "Usuário criado com sucesso, lista atualizada",
    },
    {
        "id": "T-007",
        "modulo": "USUÁRIOS",
        "descricao": "Admin cria usuário com email duplicado",
        "pre_condicoes": "Email já existe em BD",
        "passos": "1. Preencher formulário com email existente\n2. Clicar 'Salvar'",
        "resultado_esperado": "Erro 'Este email já está cadastrado'",
    },
    {
        "id": "T-008",
        "modulo": "USUÁRIOS",
        "descricao": "Admin edita usuário",
        "pre_condicoes": "Admin em /admin/users/{id}/edit",
        "passos": "1. Alterar name, role ou telefone\n2. Clicar 'Atualizar'",
        "resultado_esperado": "Usuário atualizado, toast de sucesso",
    },
    {
        "id": "T-009",
        "modulo": "USUÁRIOS",
        "descricao": "Admin desativa usuário",
        "pre_condicoes": "Usuário com isActive=true",
        "passos": "1. Clicar toggle 'Desativar' na linha do usuário\n2. Confirmar",
        "resultado_esperado": "isActive muda para false, usuário não consegue login",
    },
    {
        "id": "T-010",
        "modulo": "USUÁRIOS",
        "descricao": "Admin ativa usuário",
        "pre_condicoes": "Usuário com isActive=false",
        "passos": "1. Clicar toggle 'Ativar' na linha do usuário\n2. Confirmar",
        "resultado_esperado": "isActive muda para true, usuário consegue fazer login",
    },
    {
        "id": "T-011",
        "modulo": "USUÁRIOS",
        "descricao": "Admin altera role de Requester para Buyer",
        "pre_condicoes": "Usuário com role=requester",
        "passos": "1. Clicar editar usuário\n2. Alterar role para 'buyer'\n3. Clicar 'Atualizar'",
        "resultado_esperado": "Role alterado, usuário vê novo dashboard",
    },
    {
        "id": "T-012",
        "modulo": "USUÁRIOS",
        "descricao": "Requester edita próprio perfil",
        "pre_condicoes": "Requester em /profile/edit",
        "passos": "1. Alterar name, email ou telefone\n2. Clicar 'Salvar alterações'",
        "resultado_esperado": "Dados atualizados, toast sucesso",
    },
    {
        "id": "T-013",
        "modulo": "USUÁRIOS",
        "descricao": "Requester tenta alterar role (negado)",
        "pre_condicoes": "Requester em /profile/edit",
        "passos": "1. Tentar acessar diretamente /admin/users/{id}/edit",
        "resultado_esperado": "Erro 403 Não Autorizado",
    },
    # SOLICITAÇÕES - CRIAR (7)
    {
        "id": "T-014",
        "modulo": "SOLICITAÇÕES",
        "descricao": "Requester cria solicitação em rascunho",
        "pre_condicoes": "Requester em /requests/create",
        "passos": "1. Preencher título, centro custo, urgência\n2. Adicionar item (quantidade, unidade)\n3. Salvar rascunho",
        "resultado_esperado": "Rascunho criado com código SC-0001, status=draft",
    },
    {
        "id": "T-015",
        "modulo": "SOLICITAÇÕES",
        "descricao": "Requester adiciona múltiplos itens",
        "pre_condicoes": "Formulário aberto com Alpine.js",
        "passos": "1. Clicar '+ Adicionar Item' 3x\n2. Preencher dados em cada linha\n3. Salvar",
        "resultado_esperado": "3 SupplyRequestItems criados",
    },
    {
        "id": "T-016",
        "modulo": "SOLICITAÇÕES",
        "descricao": "Requester faz upload de anexo por item",
        "pre_condicoes": "Formulário de criação aberto",
        "passos": "1. Clique em campo 'Anexo' de item\n2. Selecionar arquivo PDF (< 10 MB)\n3. Salvar",
        "resultado_esperado": "ItemAttachment criado, arquivo em storage/app/private/",
    },
    {
        "id": "T-017",
        "modulo": "SOLICITAÇÕES",
        "descricao": "Requester tenta upload com arquivo > 10 MB",
        "pre_condicoes": "Arquivo grande selecionado",
        "passos": "1. Selecionar arquivo > 10 MB\n2. Clicar 'Salvar'",
        "resultado_esperado": "Erro 'Arquivo deve ter no máximo 10 MB'",
    },
    {
        "id": "T-018",
        "modulo": "SOLICITAÇÕES",
        "descricao": "Requester faz upload com tipo inválido",
        "pre_condicoes": "Arquivo .exe selecionado",
        "passos": "1. Selecionar arquivo .exe\n2. Clicar 'Salvar'",
        "resultado_esperado": "Erro 'Tipo de arquivo não permitido'",
    },
    {
        "id": "T-019",
        "modulo": "SOLICITAÇÕES",
        "descricao": "Requester edita rascunho",
        "pre_condicoes": "Solicitação em status=draft",
        "passos": "1. Acessar /requests/{id}/edit\n2. Alterar título e itens\n3. Clicar 'Atualizar'",
        "resultado_esperado": "Rascunho atualizado",
    },
    {
        "id": "T-020",
        "modulo": "SOLICITAÇÕES",
        "descricao": "Requester deleta rascunho",
        "pre_condicoes": "Solicitação em status=draft",
        "passos": "1. Clicar 'Deletar' em /requests/{id}\n2. Confirmar",
        "resultado_esperado": "Solicitação deletada, redireciona para lista",
    },
    # SOLICITAÇÕES - ENVIAR (3)
    {
        "id": "T-021",
        "modulo": "SOLICITAÇÕES",
        "descricao": "Requester envia rascunho para aprovação",
        "pre_condicoes": "Solicitação em status=draft com itens",
        "passos": "1. Clicar botão 'Enviar Solicitação'\n2. Confirmar",
        "resultado_esperado": "Status muda draft → pending, emails na fila",
    },
    {
        "id": "T-022",
        "modulo": "SOLICITAÇÕES",
        "descricao": "Notificações enfileiradas corretamente",
        "pre_condicoes": "Solicitação enviada, queue:work rodando",
        "passos": "1. Aguardar processamento da fila\n2. Verificar email do requester e buyers",
        "resultado_esperado": "2 emails recebidos (RequestSubmittedRequesterMail, RequestSubmittedBuyerMail)",
    },
    {
        "id": "T-023",
        "modulo": "SOLICITAÇÕES",
        "descricao": "Tentar enviar solicitação sem itens",
        "pre_condicoes": "Solicitação vazia",
        "passos": "1. Clicar 'Enviar' sem adicionar items",
        "resultado_esperado": "Erro 'Adicione pelo menos 1 item'",
    },
    # SOLICITAÇÕES - PROCESSAMENTO BUYER (10)
    {
        "id": "T-024",
        "modulo": "SOLICITAÇÕES - BUYER",
        "descricao": "Buyer visualiza solicitação pendente",
        "pre_condicoes": "Solicitação em status=pending",
        "passos": "1. Buyer acessa /requests\n2. Clica na solicitação",
        "resultado_esperado": "Página de detalhe mostra badge 'Pendente'",
    },
    {
        "id": "T-025",
        "modulo": "SOLICITAÇÕES - BUYER",
        "descricao": "Requester não vê solicitação de outro requester",
        "pre_condicoes": "2 requesters com solicitações diferentes",
        "passos": "1. Requester A acessa /requests\n2. Procura solicitação de Requester B",
        "resultado_esperado": "Não aparece (filtra por user_id)",
    },
    {
        "id": "T-026",
        "modulo": "SOLICITAÇÕES - BUYER",
        "descricao": "Buyer avança status pending → inProgress",
        "pre_condicoes": "Solicitação em status=pending",
        "passos": "1. Buyer clica 'Iniciar Atendimento'\n2. Confirma",
        "resultado_esperado": "Status muda para 'Em Andamento', RequestStatusHistory registrada",
    },
    {
        "id": "T-027",
        "modulo": "SOLICITAÇÕES - BUYER",
        "descricao": "Buyer vincula Supplier ao item",
        "pre_condicoes": "Item em status=pending",
        "passos": "1. Clicar no item\n2. Selecionar Supplier dropdown\n3. Salvar",
        "resultado_esperado": "supplier_id preenchido, item exibe nome do supplier",
    },
    {
        "id": "T-028",
        "modulo": "SOLICITAÇÕES - BUYER",
        "descricao": "Buyer avança item pending → quoting (com order_number)",
        "pre_condicoes": "Item com supplier vinculado",
        "passos": "1. Clicar 'Avançar' item\n2. Preencher order_number (ex: PC-0001)\n3. Salvar",
        "resultado_esperado": "Item status muda para 'Em Cotação', order_number salvo",
    },
    {
        "id": "T-029",
        "modulo": "SOLICITAÇÕES - BUYER",
        "descricao": "Buyer tenta avançar item sem order_number",
        "pre_condicoes": "Item em pending, sem order_number",
        "passos": "1. Clicar 'Avançar' item\n2. Deixar order_number vazio\n3. Salvar",
        "resultado_esperado": "Erro 'Order number é obrigatório'",
    },
    {
        "id": "T-030",
        "modulo": "SOLICITAÇÕES - BUYER",
        "descricao": "Buyer avança item quoting → awaitingPayment → awaitingDelivery",
        "pre_condicoes": "Item em quoting",
        "passos": "1. Clicar 'Avançar' (x2) para chegar em 'Aguardando Entrega'\n2. Confirmar",
        "resultado_esperado": "Email AwaitingDeliveryMail enfileirado para requester",
    },
    {
        "id": "T-031",
        "modulo": "SOLICITAÇÕES - BUYER",
        "descricao": "Buyer registra entrega parcial",
        "pre_condicoes": "Item em awaitingDelivery, quantity=10",
        "passos": "1. Clicar 'Registrar Entrega'\n2. Preencher delivered_quantity=5\n3. Adicionar notas\n4. Salvar",
        "resultado_esperado": "ItemDelivery criado, delivered_quantity=5, item ainda em awaitingDelivery",
    },
    {
        "id": "T-032",
        "modulo": "SOLICITAÇÕES - BUYER",
        "descricao": "Buyer registra entrega total",
        "pre_condicoes": "Item com delivered_quantity=5, quantity=10",
        "passos": "1. Clicar 'Registrar Entrega'\n2. Preencher delivered_quantity=5\n3. Salvar",
        "resultado_esperado": "ItemDelivery criado, total=10, item muda para 'Recebido'",
    },
    {
        "id": "T-033",
        "modulo": "SOLICITAÇÕES - BUYER",
        "descricao": "Solicitação auto-avança para completed",
        "pre_condicoes": "Todos os itens em status received ou cancelled",
        "passos": "1. Sistema detecta condição\n2. Auto-transição",
        "resultado_esperado": "Solicitação muda para 'Concluído', RequestCompletedMail enfileirada",
    },
    # SOLICITAÇÕES - CANCELAMENTO (6)
    {
        "id": "T-034",
        "modulo": "SOLICITAÇÕES - CANCELAMENTO",
        "descricao": "Requester solicita cancelamento de solicitação",
        "pre_condicoes": "Solicitação em status=pending",
        "passos": "1. Clicar 'Solicitar Cancelamento'\n2. Preencher motivo\n3. Confirmar",
        "resultado_esperado": "Status muda para 'Cancelamento Solicitado', emails enfileirados",
    },
    {
        "id": "T-035",
        "modulo": "SOLICITAÇÕES - CANCELAMENTO",
        "descricao": "Buyer aprova cancelamento de solicitação",
        "pre_condicoes": "Solicitação em status=cancelRequested",
        "passos": "1. Clicar 'Aprovar Cancelamento'\n2. Confirmar",
        "resultado_esperado": "Status muda para 'Cancelado', RequestCancelledMail enfileirada",
    },
    {
        "id": "T-036",
        "modulo": "SOLICITAÇÕES - CANCELAMENTO",
        "descricao": "Buyer recusa cancelamento de solicitação",
        "pre_condicoes": "Solicitação em status=cancelRequested, previous_status=pending",
        "passos": "1. Clicar 'Recusar Cancelamento'\n2. Confirmar",
        "resultado_esperado": "Status volta para previous_status (pending)",
    },
    {
        "id": "T-037",
        "modulo": "SOLICITAÇÕES - CANCELAMENTO",
        "descricao": "Admin cancela solicitação direto",
        "pre_condicoes": "Admin, solicitação em status=pending",
        "passos": "1. Clicar 'Cancelar' (ação admin)\n2. Preencher motivo\n3. Confirmar",
        "resultado_esperado": "Status muda direto para 'Cancelado', itens não finalizados → cancelado",
    },
    {
        "id": "T-038",
        "modulo": "SOLICITAÇÕES - CANCELAMENTO",
        "descricao": "Requester solicita cancelamento de item",
        "pre_condicoes": "Item em status=pending ou quoting",
        "passos": "1. Clicar 'Solicitar Cancelamento' no item\n2. Confirmar",
        "resultado_esperado": "Item status muda para 'Cancelamento Solicitado'",
    },
    {
        "id": "T-039",
        "modulo": "SOLICITAÇÕES - CANCELAMENTO",
        "descricao": "Buyer aprova cancelamento de item",
        "pre_condicoes": "Item em status=cancelRequested",
        "passos": "1. Clicar 'Aprovar' no item\n2. Confirmar",
        "resultado_esperado": "Item status muda para 'Cancelado'",
    },
    # RELATÓRIOS (5)
    {
        "id": "T-040",
        "modulo": "RELATÓRIOS",
        "descricao": "Buyer exporta solicitações em Excel",
        "pre_condicoes": "Múltiplas solicitações criadas",
        "passos": "1. Acessar /requests\n2. Clicar 'Excel'",
        "resultado_esperado": "Download de relatorio-celesta-YYYY-MM-DD.xlsx com dados",
    },
    {
        "id": "T-041",
        "modulo": "RELATÓRIOS",
        "descricao": "Buyer exporta com filtros aplicados",
        "pre_condicoes": "Filtros de status e período setados",
        "passos": "1. Aplicar filtros (status=completed, período)\n2. Clicar 'Excel'",
        "resultado_esperado": "Excel com apenas solicitações que conferem filtros",
    },
    {
        "id": "T-042",
        "modulo": "RELATÓRIOS",
        "descricao": "Buyer exporta solicitações em PDF",
        "pre_condicoes": "Solicitações listadas",
        "passos": "1. Acessar /requests\n2. Clicar 'PDF'",
        "resultado_esperado": "Download de relatório em PDF com formatação profissional",
    },
    {
        "id": "T-043",
        "modulo": "RELATÓRIOS",
        "descricao": "Requester exporta PDF de uma solicitação",
        "pre_condicoes": "Solicitação em /requests/{id}",
        "passos": "1. Clicar 'Exportar PDF' no botão",
        "resultado_esperado": "PDF individual com detalhe da solicitação",
    },
    {
        "id": "T-044",
        "modulo": "RELATÓRIOS",
        "descricao": "Rascunhos excluídos de exportação",
        "pre_condicoes": "Solicitações rascunho e pendente",
        "passos": "1. Exportar Excel/PDF",
        "resultado_esperado": "Apenas solicitação pendente aparece",
    },
    # PERMISSÕES (8)
    {
        "id": "T-045",
        "modulo": "PERMISSÕES",
        "descricao": "Requester tenta acessar admin/users",
        "pre_condicoes": "Requester autenticado",
        "passos": "1. Acessar direto /admin/users",
        "resultado_esperado": "Erro 403 Não Autorizado",
    },
    {
        "id": "T-046",
        "modulo": "PERMISSÕES",
        "descricao": "Buyer tenta avançar status item (sem permission)",
        "pre_condicoes": "Buyer tenta via URL manipulada",
        "passos": "1. POST /requests/{id}/items/{id}/status sem permission",
        "resultado_esperado": "Erro 403 Não Autorizado",
    },
    {
        "id": "T-047",
        "modulo": "PERMISSÕES",
        "descricao": "Requester vê apenas solicitações próprias",
        "pre_condicoes": "2 requesters com solicitações diferentes",
        "passos": "1. Requester A acessa /requests",
        "resultado_esperado": "Vê apenas suas solicitações (não as de Requester B)",
    },
    {
        "id": "T-048",
        "modulo": "PERMISSÕES",
        "descricao": "Buyer vê todas solicitações não-rascunho",
        "pre_condicoes": "Buyer autenticado, múltiplas solicitações",
        "passos": "1. Buyer acessa /requests",
        "resultado_esperado": "Vê todas as solicitações em status != draft",
    },
    {
        "id": "T-049",
        "modulo": "PERMISSÕES",
        "descricao": "Admin vê todas solicitações",
        "pre_condicoes": "Admin autenticado",
        "passos": "1. Admin acessa /requests",
        "resultado_esperado": "Vê todas solicitações incluindo rascunhos",
    },
    {
        "id": "T-050",
        "modulo": "PERMISSÕES",
        "descricao": "Usuário inativo não consegue acessar nada",
        "pre_condicoes": "Usuário com isActive=false",
        "passos": "1. Fazer login\n2. Acessar qualquer rota",
        "resultado_esperado": "Erro 403 'Usuário inativo ou não autorizado'",
    },
    {
        "id": "T-051",
        "modulo": "PERMISSÕES",
        "descricao": "Requester tenta deletar solicitação pendente",
        "pre_condicoes": "Solicitação em status=pending",
        "passos": "1. Clicar 'Deletar'",
        "resultado_esperado": "Erro 'Não pode deletar solicitação não-rascunho'",
    },
    {
        "id": "T-052",
        "modulo": "PERMISSÕES",
        "descricao": "Buyer tenta editar solicitação (negado)",
        "pre_condicoes": "Buyer em /requests/{id}/edit",
        "passos": "1. Acessar diretamente URL de edição",
        "resultado_esperado": "Erro 403 Não Autorizado",
    },
    # DASHBOARD (4)
    {
        "id": "T-053",
        "modulo": "DASHBOARD",
        "descricao": "Requester vê cards com contagem por status",
        "pre_condicoes": "Requester com solicitações em vários status",
        "passos": "1. Requester acessa /dashboard",
        "resultado_esperado": "Cards mostram: 2 Pendentes, 1 Em Andamento, 1 Concluído",
    },
    {
        "id": "T-054",
        "modulo": "DASHBOARD",
        "descricao": "Buyer vê card 'Pendentes de Ação' (azul)",
        "pre_condicoes": "3 solicitações em status=pending",
        "passos": "1. Buyer acessa /dashboard",
        "resultado_esperado": "Card azul mostra '3 Pendentes de Ação'",
    },
    {
        "id": "T-055",
        "modulo": "DASHBOARD",
        "descricao": "Admin vê gráfico de 12 meses",
        "pre_condicoes": "Admin no dashboard",
        "passos": "1. Admin acessa /dashboard",
        "resultado_esperado": "Gráfico Chart.js mostra solicitações por mês",
    },
    {
        "id": "T-056",
        "modulo": "DASHBOARD",
        "descricao": "Dashboard carrega corretamente após logout/login",
        "pre_condicoes": "Usuário fez logout",
        "passos": "1. Fazer login\n2. Acessar /dashboard",
        "resultado_esperado": "Dashboard renderiza com dados corretos do usuário",
    },
]

def criar_documento():
    """Cria documento Word com todos os cenários de teste"""

    doc = Document()

    # Capa
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("DOCUMENTO DE TESTES")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run("CelestaSupply - Sistema de Gestão de Solicitações de Compra")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(64, 64, 64)

    doc.add_paragraph()

    # Informações
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"Data: {datetime.now().strftime('%d/%m/%Y')}\n")
    run.font.size = Pt(11)
    run = info.add_run(f"Total de Cenários: {len(TESTE_SCENARIOS)}\n")
    run.font.size = Pt(11)

    doc.add_page_break()

    # Índice
    doc.add_heading("ÍNDICE", level=1)

    modulos = {}
    for teste in TESTE_SCENARIOS:
        mod = teste["modulo"]
        if mod not in modulos:
            modulos[mod] = 0
        modulos[mod] += 1

    for modulo, count in modulos.items():
        doc.add_paragraph(f"{modulo} ({count} cenários)", style='List Bullet')

    doc.add_page_break()

    # Cenários por módulo
    modulo_atual = None

    for teste in TESTE_SCENARIOS:
        if teste["modulo"] != modulo_atual:
            modulo_atual = teste["modulo"]
            doc.add_heading(modulo_atual, level=1)

        # Cabeçalho do teste
        heading = doc.add_paragraph()
        run = heading.add_run(f"{teste['id']}: {teste['descricao']}")
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 102, 204)

        # Tabela com detalhes
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'

        celulas = [
            ("Pré-condições", teste["pre_condicoes"]),
            ("Passos", teste["passos"]),
            ("Resultado Esperado", teste["resultado_esperado"]),
            ("Resultado Atual", "[ Preencher durante testes ]"),
            ("Status", "[ ] Passou  [ ] Falhou"),
            ("Observações", ""),
        ]

        for i, (label, valor) in enumerate(celulas):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = valor

            # Negrito na primeira coluna
            for paragraph in table.rows[i].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        doc.add_paragraph()  # Espaço

    # Resumo final
    doc.add_page_break()
    doc.add_heading("RESUMO DE TESTES", level=1)

    summary_table = doc.add_table(rows=len(modulos) + 1, cols=3)
    summary_table.style = 'Light Grid Accent 1'

    # Header
    summary_table.rows[0].cells[0].text = "Módulo"
    summary_table.rows[0].cells[1].text = "Quantidade"
    summary_table.rows[0].cells[2].text = "Passaram"

    for i, (modulo, count) in enumerate(modulos.items(), 1):
        summary_table.rows[i].cells[0].text = modulo
        summary_table.rows[i].cells[1].text = str(count)
        summary_table.rows[i].cells[2].text = "[ Preencher ]"

    doc.add_paragraph()
    doc.add_paragraph("Gerado automaticamente pelo script generate_test_scenarios.py")

    # Salvar
    output = "test-scenarios.docx"
    doc.save(output)
    print(f"✅ Documento gerado: {output}")
    print(f"📋 Total de cenários: {len(TESTE_SCENARIOS)}")
    print(f"📊 Módulos testados: {len(modulos)}")

if __name__ == "__main__":
    try:
        criar_documento()
    except ImportError:
        print("❌ Erro: python-docx não instalado")
        print("Instale com: pip install python-docx")
    except Exception as e:
        print(f"❌ Erro ao gerar documento: {e}")
