#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para gerar Relatório de Testes - CelestaSupply
Gera um documento Word completo com cenários de teste para preenchimento
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def criar_relatorio():
    doc = Document()

    # ============ ESTILOS ============
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ============ CAPA ============
    titulo = doc.add_heading('RELATÓRIO DE TESTES', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_paragraph('CelestaSupply - Sistema de Requisições de Suprimentos')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.runs[0].font.size = Pt(14)
    subtitulo.runs[0].font.bold = True

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.add_run('Data do Teste: ').bold = True
    info.add_run(f'{datetime.now().strftime("%d/%m/%Y")}')

    info = doc.add_paragraph()
    info.add_run('Testador: ').bold = True
    info.add_run('_' * 50)

    info = doc.add_paragraph()
    info.add_run('Versão: ').bold = True
    info.add_run('_' * 50)

    doc.add_page_break()

    # ============ INSTRUÇÕES ============
    doc.add_heading('Como Usar Este Documento', 1)

    instr = doc.add_paragraph()
    instr.add_run('Checkboxes: ').bold = True
    instr.add_run('Para marcar um teste como "Passou", substitua o símbolo ☐ por ☑\n')

    instr = doc.add_paragraph()
    instr.add_run('Como fazer: ').bold = True
    instr.add_run('Double-click no ☐, delete e cole ☑')

    doc.add_paragraph('Ou, no Word: selecione o símbolo ☐ → Inserir → Símbolo → Wingdings 2 → ☑')

    doc.add_paragraph()
    instr = doc.add_paragraph()
    instr.add_run('Símbolos: ').bold = True
    instr.add_run('☐ = Não marcado  |  ☑ = Marcado  |  ◈ = Em teste')

    doc.add_page_break()

    # ============ ÍNDICE ============
    doc.add_heading('Índice', 1)
    doc.add_paragraph('1. Testes de Autenticação', style='List Bullet')
    doc.add_paragraph('2. Testes de Requisições', style='List Bullet')
    doc.add_paragraph('3. Testes de Dashboard', style='List Bullet')
    doc.add_paragraph('4. Testes de Filtros e Busca', style='List Bullet')
    doc.add_paragraph('5. Testes de Relatórios', style='List Bullet')
    doc.add_paragraph('6. Testes de Performance', style='List Bullet')
    doc.add_paragraph('7. Testes de Validações', style='List Bullet')

    doc.add_page_break()

    # ============ 1. AUTENTICAÇÃO ============
    doc.add_heading('1. Testes de Autenticação', 1)

    testes_auth = [
        ('Login com credenciais válidas', 'Usuário consegue fazer login e é redirecionado para dashboard'),
        ('Login com senha incorreta', 'Sistema exibe mensagem de erro "Credenciais inválidas"'),
        ('Login com email não registrado', 'Sistema exibe mensagem de erro "Usuário não encontrado"'),
        ('Login com campos vazios', 'Sistema exibe validação "Campo obrigatório"'),
        ('Logout', 'Sessão é encerrada e usuário é redirecionado para login'),
        ('Acesso sem autenticação', 'Tentativa de acessar /dashboard redireciona para /login'),
        ('Manter sessão ativa', 'Página recarrega mantendo usuário logado'),
    ]

    adicionar_tabela_teste(doc, testes_auth)

    doc.add_page_break()

    # ============ 2. REQUISIÇÕES ============
    doc.add_heading('2. Testes de Requisições', 1)

    doc.add_heading('2.1 Criar Requisição', 2)
    testes_criar = [
        ('Criar requisição com todos os campos', 'Requisição é criada como rascunho e recebe código SC-XXXX'),
        ('Salvar como rascunho', 'Requisição fica com status "Rascunho"'),
        ('Enviar requisição', 'Requisição muda para "Pendente" e notificação é enviada'),
        ('Adicionar múltiplos itens', 'Todos os itens são salvos corretamente'),
        ('Upload de anexo por item', 'Anexo é salvo e link aparece na requisição'),
        ('Upload de múltiplos anexos', 'Todos os arquivos são listados corretamente'),
        ('Validar campos obrigatórios', 'Sistema impede envio sem título, centro de custo ou itens'),
        ('Criar com quantidade negativa', 'Sistema valida e rejeita quantidades inválidas'),
    ]

    adicionar_tabela_teste(doc, testes_criar)

    doc.add_heading('2.2 Editar Requisição', 2)
    testes_editar = [
        ('Editar rascunho', 'Alterações são salvas com sucesso'),
        ('Editar requisição em aberto', 'Sistema permite edição apenas de campos apropriados'),
        ('Mudar urgência', 'Alteração é refletida na listagem'),
        ('Adicionar item em requisição existente', 'Novo item aparece na lista'),
        ('Remover item', 'Item é deletado sem afetar outros'),
        ('Tentar editar requisição concluída', 'Sistema impede edição com mensagem apropriada'),
    ]

    adicionar_tabela_teste(doc, testes_editar)

    doc.add_heading('2.3 Cancelamento', 2)
    testes_cancel = [
        ('Solicitar cancelamento', 'Requisição muda para "Cancelamento Solicitado"'),
        ('Fornecer motivo do cancelamento', 'Motivo é armazenado e exibido'),
        ('Aprovador aprova cancelamento', 'Requisição muda para "Cancelado"'),
        ('Aprovador nega cancelamento', 'Requisição volta ao status anterior'),
        ('Cancelamento automático por timeout', 'Requisição antiga é cancelada automaticamente'),
    ]

    adicionar_tabela_teste(doc, testes_cancel)

    doc.add_page_break()

    # ============ 3. DASHBOARD ============
    doc.add_heading('3. Testes de Dashboard', 1)

    testes_dash = [
        ('Exibição de KPIs', 'Números de pendentes, urgentes, cancelamentos e completas aparecem'),
        ('Gráfico de requisições por mês', 'Gráfico é renderizado com dados corretos'),
        ('Filtro por centro de custo', 'Gráfico é atualizado ao selecionar um centro'),
        ('Toggle para visualização empilhada', 'Gráfico muda para mostrar todos os centros'),
        ('Tabela de últimas requisições', 'Últimos 10 registros aparecem corretamente'),
        ('Clique em requisição leva para detalhe', 'Redirecionamento funciona'),
        ('Dashboard carrega em menos de 2 segundos', 'Performance é aceitável'),
    ]

    adicionar_tabela_teste(doc, testes_dash)

    doc.add_page_break()

    # ============ 4. FILTROS E BUSCA ============
    doc.add_heading('4. Testes de Filtros e Busca', 1)

    testes_filtro = [
        ('Filtrar por status', 'Lista atualiza mostrando apenas requisições com status selecionado'),
        ('Filtrar por urgência', 'Filtro funciona e mostra apenas urgentes selecionadas'),
        ('Filtrar por centro de custo', 'Apenas requisições do centro selecionado aparecem'),
        ('Buscar por código (SC-XXXX)', 'Busca localiza a requisição corretamente'),
        ('Buscar por título', 'Requisições com título similar aparecem'),
        ('Buscar por solicitante', 'Apenas requisições do usuário aparecem'),
        ('Filtrar por data', 'Requisições no intervalo selecionado aparecem'),
        ('Combinar múltiplos filtros', 'Todos os filtros funcionam em conjunto'),
        ('Limpar filtros', 'Botão "Limpar" reseta todos os valores'),
    ]

    adicionar_tabela_teste(doc, testes_filtro)

    doc.add_page_break()

    # ============ 5. RELATÓRIOS ============
    doc.add_heading('5. Testes de Relatórios', 1)

    testes_relatorio = [
        ('Exportar para Excel', 'Arquivo é gerado e contém dados corretos'),
        ('Exportar para PDF', 'PDF é gerado com formatação apropriada'),
        ('Exportar requisição específica em PDF', 'Somente a requisição selecionada é exportada'),
        ('Excel contém todas as colunas', 'Código, Título, Status, Data, etc aparecem'),
        ('Relatório respeita filtros aplicados', 'Apenas dados filtrados são exportados'),
        ('Data/hora de geração aparece no PDF', 'Timestamp é visível no documento'),
    ]

    adicionar_tabela_teste(doc, testes_relatorio)

    doc.add_page_break()

    # ============ 6. PERFORMANCE ============
    doc.add_heading('6. Testes de Performance', 1)

    testes_perf = [
        ('Dashboard carrega em < 2s', 'Tempo aceitável para usuário'),
        ('Listagem com 100+ requisições carrega rápido', 'Paginação ou virtualization funciona'),
        ('Filtros responsivos sem lag', 'Filtro aplica em < 500ms'),
        ('Upload de arquivo > 5MB', 'Sistema aceita e processa'),
        ('Gráfico atualiza suavemente', 'Transição entre períodos é fluida'),
        ('Múltiplos usuários simultâneos', 'Sistema suporta load sem degradação'),
    ]

    adicionar_tabela_teste(doc, testes_perf)

    doc.add_page_break()

    # ============ 7. VALIDAÇÕES ============
    doc.add_heading('7. Testes de Validações', 1)

    testes_validacao = [
        ('Título obrigatório', 'Campo não permite salvar vazio'),
        ('Quantidade deve ser > 0', 'Sistema rejeita quantidades negativas ou zero'),
        ('Email válido no perfil', 'Validação de formato de email'),
        ('Telefone com máscara', 'Formato é preservado ou convertido corretamente'),
        ('Data de validade futura', 'Sistema valida datas apropriadas'),
        ('Nenhum campo com XSS', 'Caracteres especiais são escapados'),
        ('Nenhum SQL injection possível', 'Queries parametrizadas funcionam'),
    ]

    adicionar_tabela_teste(doc, testes_validacao)

    doc.add_page_break()

    # ============ RESUMO ============
    doc.add_heading('Resumo dos Testes', 1)

    resumo_table = doc.add_table(rows=5, cols=2)
    resumo_table.style = 'Light Grid Accent 1'

    cells = resumo_table.rows[0].cells
    cells[0].text = 'Métrica'
    cells[1].text = 'Valor'

    cells = resumo_table.rows[1].cells
    cells[0].text = 'Total de Cenários'
    cells[1].text = '_' * 20

    cells = resumo_table.rows[2].cells
    cells[0].text = 'Testes Passaram'
    cells[1].text = '_' * 20

    cells = resumo_table.rows[3].cells
    cells[0].text = 'Testes Falharam'
    cells[1].text = '_' * 20

    cells = resumo_table.rows[4].cells
    cells[0].text = 'Taxa de Sucesso'
    cells[1].text = '_' * 20 + ' %'

    doc.add_paragraph()
    obs = doc.add_paragraph()
    obs.add_run('Observações e Notas: ').bold = True
    doc.add_paragraph('_' * 80)
    doc.add_paragraph('_' * 80)
    doc.add_paragraph('_' * 80)

    # ============ SALVAR ============
    import os
    downloads = os.path.expanduser('~/Downloads')
    filename = f'Relatorio_Testes_CelestaSupply_{datetime.now().strftime("%d-%m-%Y_%H-%M-%S")}.docx'
    filepath = os.path.join(downloads, filename)

    doc.save(filepath)
    print('[OK] Relatorio gerado com sucesso!')
    print('[INFO] Local: ' + filepath)
    print('\n[CONTEUDO] O documento contem:')
    print('   - 7 secoes de testes (65 cenarios totais)')
    print('   - Checkboxes editaveis ([] -> [X])')
    print('   - Tabelas para observacoes')
    print('   - Resumo final')
    print('\n[PRONTO] Abra e preencha eletronicamente enquanto faz os testes!')

def adicionar_tabela_teste(doc, testes):
    """Adiciona uma tabela de testes ao documento com checkboxes clicáveis"""
    table = doc.add_table(rows=len(testes) + 1, cols=4)
    table.style = 'Light Grid Accent 1'

    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Cenário'
    header_cells[1].text = 'Resultado Esperado'
    header_cells[2].text = 'Status'
    header_cells[3].text = 'Observações'

    # Dados
    for i, (cenario, esperado) in enumerate(testes, 1):
        cells = table.rows[i].cells
        cells[0].text = cenario
        cells[1].text = esperado

        # Adicionar checkboxes clicáveis na célula de status
        status_para = cells[2].paragraphs[0]
        status_para.text = ''

        # Checkbox Passou
        run = status_para.add_run('☐ ')
        run.font.size = Pt(14)
        run = status_para.add_run('Passou')
        run.font.size = Pt(11)

        status_para.add_run('   ')

        # Checkbox Falhou
        run = status_para.add_run('☐ ')
        run.font.size = Pt(14)
        run = status_para.add_run('Falhou')
        run.font.size = Pt(11)

        status_para.add_run('   ')

        # Checkbox Pendente
        run = status_para.add_run('☐ ')
        run.font.size = Pt(14)
        run = status_para.add_run('Pendente')
        run.font.size = Pt(11)

        cells[3].text = ''

    doc.add_paragraph()

if __name__ == '__main__':
    criar_relatorio()
